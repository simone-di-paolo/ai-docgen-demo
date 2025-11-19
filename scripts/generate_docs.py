
import os
import subprocess
import requests
import json
import re
import shutil
from docx import Document
from docx.shared import Pt
from datetime import datetime

# === CONFIGS ===

API_KEY = os.getenv("GEMINI_API_KEY")
BOT_API_KEY = os.getenv("BOT_API_KEY")
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_CHAT_ID")

GEMINI_API_ENDPOINT = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"
SRC_FOLDER = "src/"
DOCS_DOTX_FOLDER = "docs/dotx/components/"
TEMPLATE_PATH = "docs/dotx/template_document.dotx"

GIT_BOT_EMAIL = "actions-bot@github.com"
GIT_BOT_NAME = "GitHub Actions Bot"

DEBUG = os.getenv("DEBUG_MODE", "false").lower() == "true"

# === TELEGRAM UTILS ===

def send_telegram_message(message):
    """Sends a message to a Telegram chat and truncates it if it's too long."""
    if not BOT_API_KEY or not TELEGRAM_CHAT_ID:
        if DEBUG:
            print("[DEBUG TELEGRAM] BOT_API_KEY or TELEGRAM_CHAT_ID not set. Skipping send.")
        return

    max_length = 4096  # Telegram message limit
    truncated_message = (message[:max_length - 4] + '...') if len(message) > max_length else message
    
    url = f"https://api.telegram.org/bot{BOT_API_KEY}/sendMessage"
    payload = {
        'chat_id': TELEGRAM_CHAT_ID,
        'text': truncated_message,
        'parse_mode': 'Markdown'
    }
    
    try:
        response = requests.post(url, json=payload, timeout=10)
        response.raise_for_status()
        if DEBUG:
            print(f"[DEBUG TELEGRAM] Message sent successfully.")
    except requests.exceptions.RequestException as e:
        # Do not block the script if Telegram is not working
        print(f"[TELEGRAM ERROR] Unable to send message: {e}")

# === GIT & FILE UTILS ===

def get_changed_files():
    """Get the list of edited file during in the last commit"""
    try:
        diff_output = subprocess.run(
            ["git", "diff", "HEAD~1", "HEAD", "--name-only", "--", SRC_FOLDER],
            capture_output=True, text=True, check=True
        ).stdout.strip()
        return diff_output.split('\\n') if diff_output else []
    except subprocess.CalledProcessError as e:
        error_msg = f"[ERROR] During git diff execution: {e}"
        print(error_msg)
        send_telegram_message(f"❌ ERROR: `git diff` execution failed. Cannot continue.")
        return []

def create_backup(filepath):
    """Create a backup file of the original documentation"""
    if os.path.exists(filepath):
        backup_path = f"{filepath}.bak"
        shutil.copy2(filepath, backup_path)
        msg = f"  -> Backup created: {backup_path}"
        print(msg)
        send_telegram_message(f"📑 Created backup of previous documentation: `{os.path.basename(backup_path)}`")


def get_component_name_from_path(filepath):
    """Extract the name of the component based on the file name"""
    return os.path.splitext(os.path.basename(filepath))[0]

# === AI DOCUMENTATION GENERATOR ===

def generate_new_documentation(component_name, file_diff):
    """Make a call to Gemini API to generate new documentation in JSON format."""
    prompt = f"""
    You are a senior software engineer responsible for creating clear technical documentation for a React project.
    Your task is to generate the documentation for the component '{component_name}' based on the code changes provided.
    The output must be a JSON object that will be used to programmatically generate a `.docx` file.

    **CODE CHANGES (GIT DIFF):**
    ```diff
    {file_diff}
    ```

    **INSTRUCTIONS:**
    1.  Analyze the git diff to understand the changes made to the React component.
    2.  Generate documentation that reflects these changes. If the diff is creating a new component, create the full documentation from scratch.
    3.  The documentation should be structured, clear, and technical. Write in English.
    4.  You **MUST** return a single valid JSON object. Do not include any text or markdown before or after the JSON object.
    5.  The JSON object should follow this structure:
        {{
            "title": "Component Title",
            "description": "A brief, one-sentence description of the component.",
            "sections": [
                {{
                    "heading": "Section Heading (e.g., 'Props', 'Functionality', 'Styling')",
                    "content": [
                        {{
                            "type": "paragraph",
                            "text": "A paragraph of text. You can include bold parts using **text**."
                        }},
                        {{
                            "type": "table",
                            "headers": ["Column 1", "Column 2", "Column 3"],
                            "rows": [
                                ["Row 1, Cell 1", "Row 1, Cell 2", "Row 1, Cell 3"],
                                ["Row 2, Cell 1", "Row 2, Cell 2", "Row 2, Cell 3"]
                            ]
                        }}
                    ]
                }}
            ]
        }}
    6.  For `paragraph` type, you can use markdown-like bolding with `**text**`. Example: "This paragraph contains **important** text."

    **JSON DOCUMENTATION FOR {component_name}:**
    """

    send_telegram_message(f"🧠 Calling the AI to generate JSON for `{component_name}`...")

    headers = {
        "Content-Type": "application/json",
        "x-goog-api-key": API_KEY
    }
    data = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(GEMINI_API_ENDPOINT, headers=headers, data=json.dumps(data), timeout=180)
        response.raise_for_status()
        
        raw_response = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
        
        # Clean the response to ensure it's a valid JSON
        json_str = re.search(r'\{.*\}', raw_response, re.DOTALL).group(0)
        json_data = json.loads(json_str)

        send_telegram_message(f"✅ AI responded for `{component_name}`.")
        return json_data
        
    except (requests.exceptions.RequestException, KeyError, IndexError, json.JSONDecodeError) as e:
        error_text = f"❌ ERROR processing AI response for `{component_name}`: {e}"
        print(error_text)
        send_telegram_message(error_text)
        return None

# === DOCX GENERATION ===

def replace_placeholders(doc, placeholders):
    """Replace placeholders in the entire document."""
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            if key in p.text:
                inline = p.runs
                # Replace strings and retain formatting
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text

def add_formatted_paragraph(paragraph, text):
    """Adds a paragraph with bold formatting based on markdown-like syntax."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)

def create_docx_from_json(component_name, json_data, template_path, output_path):
    """Creates a .docx file from JSON data using a template."""
    try:
        doc = Document(template_path)

        # Replace placeholders
        placeholders = {
            "{{COMPONENT_NAME}}": json_data.get("title", component_name),
            "{{COMPONENT_DESCRIPTION}}": json_data.get("description", ""),
            "{{DATE}}": datetime.now().strftime("%Y-%m-%d")
        }
        replace_placeholders(doc, placeholders)

        # Find the content placeholder paragraph
        content_placeholder = None
        for p in doc.paragraphs:
            if "{{CONTENT}}" in p.text:
                content_placeholder = p
                break
        
        if content_placeholder is None:
            raise ValueError("{{CONTENT}} placeholder not found in the template.")

        # Add sections from JSON
        for section in json_data.get("sections", []):
            if section.get("heading"):
                doc.add_paragraph(section["heading"], style='Heading 2')
            
            for item in section.get("content", []):
                if item.get("type") == "paragraph":
                    p = doc.add_paragraph(style='Body Text')
                    add_formatted_paragraph(p, item.get("text", ""))
                
                elif item.get("type") == "table":
                    headers = item.get("headers", [])
                    rows_data = item.get("rows", [])
                    if headers and rows_data:
                        table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr_cells[i].text = header
                        
                        for row_data in rows_data:
                            row_cells = table.add_row().cells
                            for i, cell_data in enumerate(row_data):
                                row_cells[i].text = str(cell_data)
        
        # Remove placeholder paragraph
        p_element = content_placeholder._element
        p_element.getparent().remove(p_element)

        doc.save(output_path)
        return True

    except Exception as e:
        error_msg = f"Error creating DOCX for {component_name}: {e}"
        print(error_msg)
        send_telegram_message(error_msg)
        return False


# === GIT OPS ===

def commit_and_push_changes(updated_docs):
    """Commit and send all the updated docs to the repo"""
    send_telegram_message("📦 Preparing to commit changes...")
    subprocess.run(["git", "config", "--global", "user.name", GIT_BOT_NAME])
    subprocess.run(["git", "config", "--global", "user.email", GIT_BOT_EMAIL])

    for doc_file in updated_docs:
        subprocess.run(["git", "add", doc_file])
    
    component_names = ', '.join([get_component_name_from_path(f) for f in updated_docs])
    commit_message = f"docs: :robot: Automatic .docx update for {component_names}"
    
    print(f"  -> Commit message: {commit_message}")
    send_telegram_message(f"📝 Commit message:`{commit_message}`")
    
    try:
        subprocess.run(["git", "commit", "-m", commit_message], check=True)
        subprocess.run(["git", "push"], check=True)
        success_msg = "🚀 Documentation sent successfully!"
        print(f"\\n>>> {success_msg} <<<")
        send_telegram_message(success_msg)
    except subprocess.CalledProcessError as e:
        error_msg = f"Git operation failed: {e}"
        print(error_msg)
        send_telegram_message(f"❌ {error_msg}")


# === MAIN WORKFLOW ===

if __name__ == "__main__":
    send_telegram_message("🤖 ===== START DOCX DOCUMENTATION SCRIPT =====")
    print("="*20 + " START DOCX DOCUMENTATION SCRIPT " + "="*20)
    
    if not os.path.exists(TEMPLATE_PATH):
        msg = f"❌ Template file not found at {TEMPLATE_PATH}. Exiting."
        print(msg)
        send_telegram_message(msg)
        exit()

    os.makedirs(DOCS_DOTX_FOLDER, exist_ok=True)

    print("\\n--- PHASE 1: Detecting Changes ---")
    changed_files = get_changed_files()
    if not changed_files or all(f == '' for f in changed_files):
        msg = "✅ No source files modified in the last commit. Exiting."
        print(msg)
        send_telegram_message(msg)
        send_telegram_message("🤖 ===== END OF DOCX SCRIPT =====")
        exit()
        
    msg = f"🔍 Detected modified files: `{'`, `'.join(changed_files)}`"
    print(msg)
    send_telegram_message(msg)

    print("\\n--- PHASE 2: Grouping Diffs by Component ---")
    component_diffs = {}
    for file_path in changed_files:
        if not (file_path.endswith(".jsx") or file_path.endswith(".css")):
            continue
        
        component_name = get_component_name_from_path(file_path)
        if component_name in ["App", "main"] and "components" in file_path:
             path_parts = file_path.split(os.sep)
             try:
                 components_index = path_parts.index("components")
                 if components_index + 1 < len(path_parts):
                     component_name = path_parts[components_index + 1]
             except ValueError:
                 pass

        if component_name not in component_diffs:
            component_diffs[component_name] = []
        
        file_diff_output = subprocess.run(
            ["git", "diff", "HEAD~1", "HEAD", "--", file_path],
            capture_output=True, text=True
        ).stdout
        component_diffs[component_name].append(file_diff_output)
    
    send_telegram_message(f"📊 Diffs grouped for the following components: `{'`, `'.join(component_diffs.keys())}`")
    print("Diffs grouped successfully.")

    print("\\n--- PHASE 3: Processing and Generating Documentation ---")
    updated_doc_files = []
    for component_name, diffs in component_diffs.items():
        doc_path = os.path.join(DOCS_DOTX_FOLDER, f"{component_name}.docx")
        full_diff = "\\n".join(diffs)

        print(f"\\n-> Processing Component: {component_name}")
        send_telegram_message(f"⚙️ Processing component: *{component_name}*")
        
        json_documentation = generate_new_documentation(component_name, full_diff)

        if json_documentation:
            create_backup(doc_path)
            if create_docx_from_json(component_name, json_documentation, TEMPLATE_PATH, doc_path):
                msg = f"  -> Documentation for {component_name} updated by the AI."
                print(msg)
                send_telegram_message(f"✍️ DOCX documentation for `{component_name}` updated.")
                updated_doc_files.append(doc_path)
            else:
                msg = f"  -> Failed to create DOCX for {component_name}."
                print(msg)
                send_telegram_message(f"❌ Failed to create DOCX for `{component_name}`.")
        else:
            msg = f"  -> No documentation generated for {component_name}."
            print(msg)
            send_telegram_message(f"🤷‍♂️ No documentation generated for `{component_name}`.")

    print("\\n--- PHASE 4: Finalization ---")
    if updated_doc_files:
        print("Documentation updates detected. Submitting changes...")
        commit_and_push_changes(updated_doc_files)
    else:
        msg = "✅ No documentation files were modified. No commit needed."
        print(msg)
        send_telegram_message(msg)
        
    send_telegram_message("🤖 ===== END OF DOCX SCRIPT =====")
